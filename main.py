import tiktoken
import requests

from openai import OpenAI
from bs4 import BeautifulSoup 
from docx import Document
from config import AI_TOKEN, SEARCH_API_KEY

def serper_search(query):
    api_key = SEARCH_API_KEY
    url = "https://serpapi.com/search"
    params = {
        "engine": "google",
        "q": query,
        'api_key': api_key,
        'gl' : "ru",
        'hl' : "ru",
        "sort_by": "date",
    }
    response = requests.get(url,  params=params)
    try:
        return response.json()
    except:
        print(response.text)
        return None
    
    
def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))    


def save_docs(text:str, file_name:str):
    doc = Document()
    doc.add_heading(f'Текс для {file_name}', 0)
    doc.add_paragraph(text)
    doc.save(f'{file_name}.docx')
    
    
def save_image(image_url, file_name:str):
    response = requests.get(image_url)
    if response.status_code == 200:
        with open(file_name, 'wb') as file:
            file.write(response.content)
        
        print(f"Изображение успешно сохранено в: {file_name}")
    else:
        print(f"Не удалось загрузить изображение. Статус: {response.status_code}")
    
    
    
def get_page_text(url:str):
    headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.121 Safari/537.36"
    }
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        
        page_text = soup.get_text(separator=' ', strip=True)
        return page_text
    else:
        print(f"Error: {response.status_code}")
        return '\n'


def prompt_to_blog(client: OpenAI, text: str, tagret_human: str='не указана'):
    
    chunk_size = 4096
    messages = []
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i + chunk_size]
        messages.append({'role': 'user', 'content': chunk})
        total_tokens = sum(count_tokens(msg["content"]) for msg in messages)
        if total_tokens > 20000 - chunk_size:
            break
    final_messages = {'role': 'user', 'content': f'Ты профессиональный специалист по анализу текстов.\
            Это все части текста. Проанализируй и напиши текст для блога, длиной не менее 150 и не более 300 слов.\
            Исключи рекламу, она не должна фигурировать в тексте и ты не должен на нее опираться при анализе.\
            Твоей целевой аудиторией являются - {tagret_human}. '}
    messages.append(final_messages)
    response =  client.chat.completions.create(
        model="gpt-4o",
        messages = messages,
        temperature=0.7,
        max_tokens=16384,
    )   
    return response.choices[0].message.content


def promt_to_create_request(client: OpenAI, text: str, tagret_human: str='не указана'):
    response =  client.chat.completions.create(
        model="gpt-4o-mini",
        messages = [
            {'role': 'user', 'content': 'Тебе отдают запрос, который человек хочет отправить в браузер.\
            Твоя задача проанализировать запрос и переписать его таким образом, \
            чтоб при запросе в браузере ссылки отражали все аспекты данного запроса и минимизировать рекламу по данным ссылкам.\
            Твой ответ должен содержать только этот запрос и ничего больше. '+text}
        ],
        temperature=0.7,
        max_tokens=1600,
    )
    return response.choices[0].message.content


def prompt_to_title(client: OpenAI, text: str, tagret_human: str='не указана'):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages = [
            {'role': 'user', 'content': 'У тебя будут четыре задачи, вторая будет зависеть от первой'},
            {'role': 'user', 'content': f'1) Из текста блога выдели главную мысль или лозунг.{text}'},
            {'role': 'user', 'content': '2) Ты профессиональный художник.\
            Твоя задача подробно описать изображение, которое ты бы нарисовал по мысли или лозунгу,\
                которые ты выделил в преведущей задаче'},
            {'role': 'user', 'content': '3) Из составленного тобой описание в предущем пункте,\
                составь промт для генерации изображения.' },
            {'role': 'user', 'content': '4) переведи этот промт на английский язык.  В ответ запиши только этот промт.'}
        ]
    )
    return response.choices[0].message.content


def generate_image(client: OpenAI, text: str, tagret_human: str='не указана'):
    response = client.images.generate(
        prompt =text,
        n=1,
        size="1024x1024"
    )
    return response.data[0].url

client = OpenAI(api_key=AI_TOKEN)
client_query = 'Новости, касающиеся важности раннего обучения детей арифметике, скорочтению и другим предметам.' 
tagret_human = 'Молодые мамы'
query = promt_to_create_request(client, client_query)
result = serper_search(query)

text_blog = ''
for url in result["organic_results"]:
    text = get_page_text(url['link'])
    text_blog += '\n' + text
    
query = prompt_to_blog(client, text_blog, tagret_human)
save_docs(query, file_name='blog')
title = prompt_to_title(client, query, tagret_human)
save_docs(title, file_name='prompt_image')
image_url = generate_image(client, title, tagret_human)
save_image(image_url, file_name='image.png')



